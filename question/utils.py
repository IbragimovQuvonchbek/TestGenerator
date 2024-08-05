import random
from .models import Question, Option
from docx import Document
from io import BytesIO
from django.http import HttpResponse
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

small_width = Inches(2 / 1)
very_small_width = Inches(2 / 9)


def set_font_size(doc, size_pt):
    """Set font size for all runs in all paragraphs of the document."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)


def create_docx_file():
    # Load the template DOCX file
    template_path = 'template.docx'
    doc = Document(template_path)

    # Ensure there is at least one section
    if not doc.sections:
        doc.add_section()

    # Set number of columns for the first section
    section = doc.sections[0]
    section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')  # Set number of columns

    # Optional: Adjust column width and spacing (default values can be adjusted as needed)
    cols = section._sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:space'), '720')  # 720 twips = 0.5 inch spacing between columns

    # Add your content to the template
    questions = list(Question.objects.all())
    random.shuffle(questions)
    index_question = 1

    for item in questions:
        question = item
        is_indexed = False
        if question.question_text:
            is_indexed = True
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"{index_question}. ")
            run.bold = True
            paragraph.add_run(question.question_text)
        if question.question_picture:
            if is_indexed:
                doc.add_paragraph()  # Add a blank line before the picture
            doc.add_picture(question.question_picture, width=small_width)

        options = list(Option.objects.filter(question=question))
        random.shuffle(options)
        index_option = 0

        option_texts = []  # To store options temporarily

        for option in options:
            option_text = f"{chr(65 + index_option)}) {option.option_text}    "
            image_stream = None
            if option.option_picture:
                image_data = option.option_picture.file.read()
                image_stream = BytesIO(image_data)  # Correctly create a BytesIO object from bytes
                option_text += ' '  # Space before the image
            option_texts.append((option_text, image_stream))
            index_option += 1

        # Write options to paragraphs, two options per line
        for i in range(0, len(option_texts), 2):
            paragraph = doc.add_paragraph()
            if i < len(option_texts):
                run = paragraph.add_run(option_texts[i][0])
                run.bold = False
                if option_texts[i][1]:
                    paragraph.add_run().add_picture(option_texts[i][1], width=very_small_width)
            if i + 1 < len(option_texts):
                run = paragraph.add_run(option_texts[i + 1][0])
                run.bold = False
                if option_texts[i + 1][1]:
                    paragraph.add_run().add_picture(option_texts[i + 1][1], width=very_small_width)

        index_question += 1

    # Set font size to 9 points for the entire document
    set_font_size(doc, 9)

    # Save the modified document to a BytesIO object
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    # Create an HttpResponse to send the file
    response = HttpResponse(doc_io,
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=randomized_questions.docx'
    return response
